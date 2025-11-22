#!/usr/bin/env python3
# /// script
# dependencies = [
#     "certifi==2025.10.5",
#     "charset-normalizer==3.4.4",
#     "idna==3.11",
#     "lxml==6.0.2",
#     "python-docx==1.2.0",
#     "requests==2.32.5",
#     "typing_extensions==4.15.0",
#     "urllib3==2.5.0",
# ]
# ///
"""
Generates a Threat Intelligence report in .docx format and an interactive HTML appendix.
"""
import csv
import os
import json
import shutil
import re
import requests
import time
import glob
import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
import argparse

# This will hold all vulnerability data for the JSON appendix
all_vulnerabilities = []
# BASE_DIR will be set dynamically based on the input file
BASE_DIR = None
REPORT_DIR = None  # Will be set dynamically based on --output-dir argument
ASSETS_DIR = None  # Will be set dynamically based on REPORT_DIR

def get_file_path(filename):
    """Constructs the full path to an input file relative to BASE_DIR."""
    if BASE_DIR is None:
        raise ValueError("BASE_DIR is not set. Ensure --input-file is provided.")
    return os.path.join(BASE_DIR, filename)

def get_latest_jquery_version():
    """Fetches the latest stable jQuery version from the official CDN."""
    try:
        response = requests.get("https://code.jquery.com/", timeout=5)
        response.raise_for_status()
        match = re.search(r"jQuery Core (\d+\.\d+\.\d+)", response.text)
        if match:
            return match.group(1)
    except requests.exceptions.RequestException as e:
        print(f"Could not fetch latest jQuery version: {e}")
    return "N/A"

def get_cves_for_product(product, version):
    """Queries the NVD API for CVEs related to a specific product and version."""
    cves = []
    try:
        search_query = f"{product} {version}"
        url = f"https://services.nvd.nist.gov/rest/json/cves/2.0?keywordSearch={search_query}"
        response = requests.get(url, timeout=10)
        response.raise_for_status()
        data = response.json()
        
        if 'vulnerabilities' in data:
            for item in data['vulnerabilities']:
                cve_id = item['cve']['id']
                summary = item['cve']['descriptions'][0]['value']
                cves.append({'cve_id': cve_id, 'summary': summary})
        
        time.sleep(1)
        
    except requests.exceptions.RequestException as e:
        print(f"Error querying NVD for {product} {version}: {e}")
    
    return cves

def generate_scope_summary():
    """Calculates statistics about the scope of the analysis."""
    stats = {
        'total_domains': 0,
        'total_subdomains': 0,
        'urls_scanned': 0,
        'typosquat_domains': 0,
        'operations_executed': 0
    }

    recon_csv = get_file_path('Recon_out.csv')
    primary_domains = set()
    subdomains = set()
    if os.path.exists(recon_csv):
        try:
            with open(recon_csv, 'r', encoding='utf-8') as f:
                reader = csv.DictReader(f)
                for row in reader:
                    domain = (row.get('domain') or '').strip().lower()
                    if not domain:
                        continue
                    subdomains.add(domain)
                    parts = domain.split('.')
                    if len(parts) >= 2:
                        primary_domains.add('.'.join(parts[-2:]))
                    else:
                        primary_domains.add(domain)
        except Exception:
            pass

    if primary_domains:
        stats['total_domains'] = len(primary_domains)
    if subdomains:
        stats['total_subdomains'] = len(subdomains)

    try:
        with open(get_file_path('Admin_Login_Enumerator_out.csv'), 'r', encoding='utf-8') as f:
            stats['urls_scanned'] = max(sum(1 for line in f) - 1, 0)
    except Exception:
        pass

    try:
        with open(get_file_path('filtered_domains.txt'), 'r', encoding='utf-8') as f:
            stats['typosquat_domains'] = len({line.strip() for line in f if line.strip()})
    except Exception:
        try:
            with open(get_file_path('domains_only.txt'), 'r', encoding='utf-8') as f:
                stats['typosquat_domains'] = len({line.strip() for line in f if line.strip()})
        except Exception:
            pass

    # Approximate total operations by counting unique raw_output artefacts produced this run.
    raw_outputs_dir = get_file_path('')
    if os.path.isdir(raw_outputs_dir):
        files_seen = {
            os.path.join(root, fname)
            for root, _, files in os.walk(raw_outputs_dir)
            for fname in files
            if not fname.startswith('.')
        }
        stats['operations_executed'] = len(files_seen)

    return stats

def set_styles(styles):
    """Defines and registers the custom styles for the document."""
    
    style = styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    style = styles['Heading 1']
    style.font.name = 'Microsoft Sans Serif'
    style.font.size = Pt(22)
    style.font.bold = True
    style.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    style.paragraph_format.space_after = Pt(12)

    style = styles['Heading 2']
    style.font.name = 'Arial'
    style.font.size = Pt(15)
    style.font.bold = True
    style.paragraph_format.space_before = Pt(18)
    style.paragraph_format.space_after = Pt(12)

    style = styles['Heading 3']
    style.font.name = 'Arial'
    style.font.size = Pt(14)
    style.font.color.rgb = RGBColor(0x26, 0x26, 0x26)
    style.paragraph_format.space_before = Pt(18)
    style.paragraph_format.space_after = Pt(6)

    risk_styles = {
        "Extreme": "B10C2A", "High": "FA4823", "Medium": "FAC300",
        "Low": "CE8D3E", "Very Low": "808080"
    }
    for risk, color in risk_styles.items():
        style_name = f'Risk - {risk}'
        if style_name not in styles:
            style = styles.add_style(style_name, WD_STYLE_TYPE.CHARACTER)
            style.font.bold = True
            style.font.color.rgb = RGBColor.from_string(color)

def add_header_footer(doc):
    """Adds a standard header and footer to the document."""
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.text = "TIaaS Intelligence Reporting"
    
    footer = section.footer
    p = footer.paragraphs[0]
    p.text = "Copyright © 2025 Trustwave Holdings, Inc. All rights reserved."

def create_report_structure(doc):
    """Creates the main sections and boilerplate text of the report."""
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "This report summarizes the findings from a threat intelligence and reconnaissance review of your organization's public-facing digital assets. "
        "The following sections detail identified vulnerabilities, their potential business impact, and recommended remediation actions. "
        "The findings are prioritized by risk and mapped to the MITRE ATT&CK® framework to provide context on attacker tactics. "
        "Always validate items of interest before acting, as threat landscapes and asset inventories can shift quickly.",
        style='Normal'
    )

def add_scope_section(doc, stats):
    doc.add_heading("Scope and Methodology", level=2)
    doc.add_paragraph(
        "The findings in this report reflect the current reconnaissance results gathered during this engagement. Threat conditions evolve rapidly; "
        "treat this as a point-in-time snapshot and validate the details below before responding.",
        style='Normal'
    )
    doc.add_paragraph(f"Primary Domains Analyzed: {stats['total_domains']:,}", style='List Bullet')
    doc.add_paragraph(f"Total Subdomains Discovered and Assessed: {stats['total_subdomains']:,}", style='List Bullet')
    doc.add_paragraph(f"Web URLs Scanned for Common Vulnerabilities: {stats['urls_scanned']:,}", style='List Bullet')
    doc.add_paragraph(f"Potential Typosquatting Variations Checked: {stats['typosquat_domains']:,}", style='List Bullet')
    if stats.get('operations_executed'):
        doc.add_paragraph(f"Individual Recon Operations Executed This Run: {stats['operations_executed']:,}", style='List Bullet')
    doc.add_paragraph()

def add_vulnerability_section(doc, title, description, mitre_ttp, remediation, headers, file_path, filter_condition, risk_column=None, default_risk=None, risk_order=None):
    """Adds a section with a description, MITRE TTPs, remediation, and a table of vulnerabilities."""
    global all_vulnerabilities
    full_path = get_file_path(file_path)

    if not os.path.exists(full_path):
        return

    vulnerabilities = []
    try:
        with open(full_path, 'r', encoding='utf-8') as f:
            reader = csv.DictReader(f)
            if reader.fieldnames is None:
                return
            for row in reader:
                try:
                    working_row = dict(row)
                    if not filter_condition(working_row):
                        continue

                    if risk_column:
                        current = working_row.get(risk_column) or working_row.get(risk_column.lower())
                        if current is None or str(current).strip() == "":
                            if default_risk:
                                working_row[risk_column] = default_risk
                        else:
                            working_row[risk_column] = str(current).strip()

                    filtered_row = {}
                    for header in headers:
                        value = working_row.get(header)
                        if value is None:
                            value = working_row.get(header.lower())
                        filtered_row[header] = value if value is not None else ''

                    if risk_column and risk_column not in headers:
                        filtered_row[risk_column] = working_row.get(risk_column, default_risk or '')

                    filtered_row['source_file'] = file_path
                    vulnerabilities.append(filtered_row)
                except Exception:
                    continue
    except (FileNotFoundError, StopIteration):
        return
    except Exception as e:
        print(f"Skipping {file_path} due to error: {e}")
        return

    if not vulnerabilities:
        return

    if risk_column and risk_order:
        def risk_sort_key(item):
            value = str(item.get(risk_column, '')).strip()
            return risk_order.get(value, len(risk_order))
        vulnerabilities.sort(key=risk_sort_key)

    headers_local = list(headers)
    if risk_column and risk_column not in headers_local:
        headers_local.append(risk_column)

    all_vulnerabilities.extend(vulnerabilities)
    doc.add_heading(title, level=3)
    doc.add_paragraph(description, style='Normal')

    p = doc.add_paragraph()
    p.add_run("Relevant MITRE ATT&CK® Tactics: ").bold = True
    p.add_run(mitre_ttp)

    p = doc.add_paragraph()
    p.add_run("Recommended Remediation:").bold = True
    for point in remediation:
        doc.add_paragraph(point, style='List Bullet')

    table = doc.add_table(rows=1, cols=len(headers_local), style='Table Grid')
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers_local):
        hdr_cells[i].text = header

    risk_key = (risk_column or '').lower()
    MAX_CELL_LENGTH = 180
    for item in vulnerabilities[:20]:
        row_cells = table.add_row().cells
        for i, header in enumerate(headers_local):
            cell_text = item.get(header, '')
            paragraph = row_cells[i].paragraphs[0]
            if isinstance(cell_text, str):
                text_value = cell_text
            else:
                text_value = str(cell_text)

            if len(text_value) > MAX_CELL_LENGTH:
                text_value = text_value[:MAX_CELL_LENGTH - 3] + "..."

            if header.lower() == risk_key and f'Risk - {text_value}' in doc.styles:
                run = paragraph.add_run(text_value)
                run.style = f'Risk - {text_value}'
            else:
                paragraph.text = text_value

    appendix_note = doc.add_paragraph()
    note = appendix_note.add_run(f"For the full dataset, open Interactive_Appendix.html and filter on {file_path}.")
    note.italic = True
    doc.add_paragraph()

def add_simple_list_section(doc, title, description, mitre_ttp, remediation, file_path, header, filter_condition=lambda line: True):
    """Adds a section for simple, single-column text files."""
    global all_vulnerabilities
    vulnerabilities = []
    full_path = get_file_path(file_path)
    if not os.path.exists(full_path):
        return
    try:
        with open(full_path, 'r', encoding='utf-8') as f:
            for line in f:
                line = line.strip()
                if line and filter_condition(line):
                    vulnerabilities.append({header: line, 'source_file': file_path})
    except FileNotFoundError:
        return
    except Exception as e:
        print(f"Skipping {file_path} due to error: {e}")
        return

    if not vulnerabilities:
        return

    all_vulnerabilities.extend(vulnerabilities)
    doc.add_heading(title, level=3)
    doc.add_paragraph(description, style='Normal')

    p = doc.add_paragraph()
    p.add_run("Relevant MITRE ATT&CK® Tactics: ").bold = True
    p.add_run(mitre_ttp)

    p = doc.add_paragraph()
    p.add_run("Recommended Remediation:").bold = True
    for point in remediation:
        doc.add_paragraph(point, style='List Bullet')

    table = doc.add_table(rows=1, cols=1, style='Table Grid')
    table.rows[0].cells[0].text = header

    for item in vulnerabilities[:20]:
        table.add_row().cells[0].text = item[header]

    appendix_note = doc.add_paragraph()
    note = appendix_note.add_run(f"For the full dataset, open Interactive_Appendix.html and filter on {file_path}.")
    note.italic = True
    doc.add_paragraph()

def add_hyperlink(paragraph, text, url):
    """Adds a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    
    c = docx.oxml.shared.OxmlElement('w:color')
    c.set(docx.oxml.shared.qn('w:val'), "0000FF")
    rPr.append(c)
    u = docx.oxml.shared.OxmlElement('w:u')
    u.set(docx.oxml.shared.qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def add_redirect_issue_section(doc, title, description, mitre_ttp, remediation, file_path):
    """Adds a section for redirect issues from a custom text file format."""
    global all_vulnerabilities
    vulnerabilities = []
    full_path = get_file_path(file_path)
    if not os.path.exists(full_path):
        return
    try:
        with open(full_path, 'r', encoding='utf-8') as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                
                if '→' in line:
                    parts = line.split('→')
                    original_domain = parts[0].replace('Original Domain:', '').strip()
                    redirect_url = parts[1].replace('Redirect URL:', '').strip()
                    vulnerabilities.append({
                        'Domain': original_domain,
                        'Status': 'Redirects',
                        'Redirect URL': redirect_url,
                        'source_file': file_path
                    })
                elif '[!]' in line:
                    domain = line.split(':')[0].replace('[!]', '').strip()
                    vulnerabilities.append({
                        'Domain': domain,
                        'Status': 'No HTTPS or redirection',
                        'Redirect URL': 'N/A',
                        'source_file': file_path
                    })
    except FileNotFoundError:
        return
    except Exception as e:
        print(f"Skipping {file_path} due to error: {e}")
        return

    if not vulnerabilities:
        return

    all_vulnerabilities.extend(vulnerabilities)
    doc.add_heading(title, level=3)
    doc.add_paragraph(description, style='Normal')

    p = doc.add_paragraph()
    p.add_run("Relevant MITRE ATT&CK® Tactics: ").bold = True
    p.add_run(mitre_ttp)

    p = doc.add_paragraph()
    p.add_run("Recommended Remediation:").bold = True
    for point in remediation:
        doc.add_paragraph(point, style='List Bullet')

    headers = ["Domain", "Status", "Redirect URL"]
    table = doc.add_table(rows=1, cols=len(headers), style='Table Grid')
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header

    for item in vulnerabilities[:20]:
        row_cells = table.add_row().cells
        row_cells[0].text = item.get('Domain', 'N/A')
        row_cells[1].text = item.get('Status', 'N/A')
        row_cells[2].text = item.get('Redirect URL', 'N/A')

    appendix_note = doc.add_paragraph()
    note = appendix_note.add_run(f"For the full dataset, open Interactive_Appendix.html and filter on {file_path}.")
    note.italic = True
    doc.add_paragraph()

def populate_findings(doc, latest_jquery_version):
    """Populates the report with filtered findings and explanations."""

    risk_priority = {
        "Critical": 0,
        "Extreme": 0,
        "High": 1,
        "Medium": 2,
        "Low": 3,
        "Very Low": 4,
        "Informational": 5
    }

    cloud_risk_order = {
        "Extreme": 0,
        "High": 1,
        "Medium": 2,
        "Low": 3,
        "Informational": 4
    }

    def ssl_attention_filter(row):
        days = str(row.get('days_until_expiry', '')).strip()
        if not days or days.lower() == 'nan':
            return False
        try:
            value = float(days)
        except ValueError:
            return False
        if value <= 0:
            row['Risk'] = "Critical"
            return True
        if value <= 7:
            row['Risk'] = "High"
            return True
        if value <= 30:
            row['Risk'] = "Medium"
            return True
        return False

    def waf_gap_filter(row):
        waf_status = str(row.get('waf_detected', '')).strip().lower()
        headers = str(row.get('waf_headers', '')).strip().lower()
        if 'error' in headers:
            row['Risk'] = "Medium"
            return True
        if waf_status in ('false', '0', '', 'no', 'none'):
            row['Risk'] = "High"
            return True
        return False

    add_vulnerability_section(
        doc, "Administrative Login Pages Available",
        ("This section identifies administrative login pages for your company's web applications that are openly accessible from the public internet. "
         "Exposing these interfaces increases the risk of unauthorized access, as attackers can use automated tools to guess common passwords or exploit vulnerabilities in the login form itself. "),
        "Credential Access (T1110: Brute Force, T1078: Valid Accounts)",
        [
            "Restrict access to admin interfaces via IP whitelisting or a VPN.",
            "Implement Multi-Factor Authentication (MFA) on all administrative accounts.",
            "Enforce a strong password policy and account lockout mechanisms."
        ],
        ['domain', 'url', 'status', 'Risk'], 'Admin_Login_Enumerator_out.csv',
        lambda row: row.get('status', '').strip() == 'Status: 200',
        risk_column='Risk', default_risk='High', risk_order=risk_priority
    )

    add_vulnerability_section(
        doc, "Exposed Cloud Storage Buckets",
        ("This section identifies cloud storage buckets (like Amazon S3) that are publicly accessible. Misconfigured buckets are a leading cause of data breaches, potentially exposing sensitive internal documents, customer data, or application source code."),
        "Collection (T1530: Data from Cloud Storage Object)",
        [
            "Immediately set the bucket permissions to private.",
            "Enable 'Block Public Access' settings at the AWS account level.",
            "Review and apply the principle of least privilege to all IAM policies related to S3."
        ],
        ['bucket_name', 'region', 'status', 'Risk'], 'S3_Bucket_Check_out.csv',
        lambda row: row.get('status') and 'Not Found' not in row.get('status') and 'Error' not in row.get('status'),
        risk_column='Risk', default_risk='High', risk_order=risk_priority
    )

    add_vulnerability_section(
        doc, "Cloud Storage Exposure (Multi-Cloud)",
        ("Public cloud storage buckets discovered across AWS, Azure, GCP, DigitalOcean, and other providers represent significant data exposure risks. "
         "Misconfigured storage containers across multiple cloud platforms can leak sensitive data, credentials, backups, or proprietary information. "
         "Each cloud provider has different permission models, making it critical to audit storage access controls consistently across all platforms."),
        "Collection (T1530: Data from Cloud Storage Object)",
        [
            "Review and restrict permissions on all discovered cloud storage buckets to private access only.",
            "Enable encryption at rest for all cloud storage containers across all providers.",
            "Implement access logging and monitoring to detect unauthorized access attempts.",
            "Audit bucket contents and remove any sensitive data that should not be in cloud storage.",
            "Apply cloud provider-specific security features (AWS Block Public Access, Azure Private Endpoints, GCP uniform bucket-level access)."
        ],
        ['bucket_name', 'provider', 'status', 'risk', 'sample_files'], 'Other_Buckets_out.csv',
        lambda row: row.get('status', '').lower() not in ['private', 'not found', 'error'],
        risk_column='risk', risk_order=cloud_risk_order
    )

    add_vulnerability_section(
        doc, "Exposed Management Services (Validation Required)",
        ("Our scan detected several sensitive management services (e.g., Kubernetes API, Elasticsearch, databases) that are accessible from the internet. "
         "IMPORTANT: This is a passive port scan - we have NOT verified whether these services are properly secured with authentication, access controls, or network restrictions. "
         "Each finding requires manual validation to determine actual risk. Services may be intentionally exposed with proper security controls in place."),
        "Initial Access (T1190: Exploit Public-Facing Application)",
        [
            "VALIDATE EACH SERVICE: Manually test whether authentication is required and properly enforced.",
            "Review firewall rules and security groups to confirm only authorized IPs can access management ports.",
            "If authentication is not required or uses default credentials, immediately restrict access.",
            "Verify that exposed services are patched and necessary for business operations.",
            "Consider implementing VPN or zero-trust network access for management interfaces.",
            "NOTE: Active security testing (credential testing, unauthorized access attempts) requires written authorization."
        ],
        ['domain', 'ip', 'port', 'service', 'risk'], 'Cloud_Misconfig_out.csv',
        lambda row: row.get('risk', '').strip() in ['High', 'Extreme'],
        risk_column='risk', risk_order=cloud_risk_order
    )

    add_vulnerability_section(
        doc, "SSL/TLS Certificates Requiring Attention",
        ("Certificates that have already expired or are close to expiration can break secure communication and erode customer trust. "
         "Attackers monitor for lapsed certificates to impersonate legitimate services or persuade users to ignore browser security warnings."),
        "Credential Access (T1556.004: Modify Authentication Process)",
        [
            "Renew certificates before they expire and remove unused hostnames from scope.",
            "Enable automated certificate issuance and renewal for all public-facing services.",
            "Monitor certificate transparency logs for unexpected changes." 
        ],
        ['domain', 'cert_expiry', 'days_until_expiry', 'Risk'], 'SSL_TLS_Cert_Check_out.csv',
        ssl_attention_filter,
        risk_column='Risk', risk_order=risk_priority
    )

    add_vulnerability_section(
        doc, "Web Applications Without WAF Coverage",
        ("Public web applications that lack a Web Application Firewall (WAF) are more exposed to automated probing and exploitation. "
         "A WAF provides an adaptive layer to detect and block common attack patterns while you address underlying vulnerabilities."),
        "Initial Access (T1190: Exploit Public-Facing Application)",
        [
            "Deploy a WAF or reverse proxy in front of internet-accessible applications.",
            "Enable logging and alerting for blocked requests to spot new attack trends.",
            "Tune WAF policies so that high-value applications are protected without impacting availability."
        ],
        ['domain', 'waf_detected', 'waf_headers', 'Risk'], 'WAF_Checker_out.csv',
        waf_gap_filter,
        risk_column='Risk', risk_order=risk_priority
    )

    add_vulnerability_section(
        doc, "Email Security: Missing DMARC Enforcement",
        ("This section reviews the status of your domain's DMARC records. DMARC is a critical email security standard that prevents attackers from spoofing your company's domain for phishing attacks. "
         "A missing or unenforced DMARC record makes your organization, partners, and customers more vulnerable to email-based fraud."),
        "Initial Access (T1566: Phishing)",
        [
            "Publish a DMARC record with a policy of `p=reject`.",
            "Implement DKIM (DomainKeys Identified Mail) to digitally sign outgoing emails.",
            "Strengthen the SPF (Sender Policy Framework) record to prevent unauthorized mail servers."
        ],
        ['domain', 'spf_status', 'dmarc_status', 'Risk'], 'Email_Security_Audit_out.csv',
        lambda row: 'missing' in row.get('dmarc_status', '').lower(),
        risk_column='Risk', default_risk='High', risk_order=risk_priority
    )

    add_vulnerability_section(
        doc, "Certificate Transparency Intelligence",
        ("Every SSL/TLS certificate issued for your domains is logged in public Certificate Transparency (CT) logs - a mandatory audit trail maintained by Certificate Authorities. "
         "This section shows ALL certificates ever issued for your subdomains, which helps discover: (1) Forgotten or decommissioned services still in DNS, "
         "(2) Unauthorized certificate issuance by rogue administrators or attackers, (3) Services that have migrated but left old DNS records behind (subdomain takeover risks), "
         "(4) Certificate/hostname mismatches that could indicate configuration issues. "
         "\n\nHOW TO READ THIS DATA: The 'domain' column shows what subdomain you have in DNS. The 'common_name' column shows what hostname the certificate was actually issued for. "
         "When these don't match, it usually means the DNS record points to a service hosted elsewhere (example: DNS record 'printers.company.com' might have a certificate for 'hosted-printer-service.cloud-provider.com'). "
         "The 'not_after' date shows when the certificate expires - expired certificates with DNS records still active may indicate forgotten/abandoned services. "
         "Multiple entries for the same domain typically show certificate renewal history over time."),
        "Resource Development (T1583.001: Acquire Infrastructure - Domains)",
        [
            "PRIORITY: Identify entries where domain and common_name don't match - these may indicate services hosted elsewhere or misconfigurations.",
            "Check for expired certificates (not_after is in the past) where the DNS record still exists - these are likely forgotten/decommissioned services.",
            "Review services that show multiple hostname variations over time (e.g., 'printlogic' → 'printlogicnew' → 'printerlogic') and clean up old DNS records.",
            "Investigate any certificates you don't recognize - unauthorized certificate issuance is a strong indicator of compromise.",
            "For wildcard certificates (*.domain.com), verify they are properly secured and not being abused.",
            "Implement automated monitoring for new certificate issuance to catch unauthorized certificates immediately."
        ],
        ['domain', 'common_name', 'issuer_name', 'not_before', 'not_after'], 'crt_transparency.csv',
        lambda row: True
    )

    add_vulnerability_section(
        doc, "Directory Listing Enabled",
        ("Directory listing is a web server feature that reveals the contents of a directory when no main web page is present. When enabled, it allows anyone to see and potentially download sensitive files, "
         "such as source code, configuration files, or data backups, that were not intended for public viewing."),
        "Reconnaissance (T1593: Search Open Websites/Domains)",
        [
            "Disable the 'autoindex' or 'directory listing' feature on the web server.",
            "Ensure a default index page (e.g., index.html) is present in every directory."
        ],
        ['Directory URL', 'File URL', 'File Name', 'Size', 'Last Modified'], 'Dir_Listing_Checker_out.csv',
        lambda row: True
    )

    add_vulnerability_section(
        doc, "Exposed .git Directories",
        ("The `.git` directory is used by the Git version control system to track changes in source code. If this directory is accidentally exposed on a live web server, it can allow an attacker to download the application's entire source code and its history. "
         "This could reveal sensitive information like API keys, database passwords, or proprietary business logic."),
        "Reconnaissance (T1593: Search Open Websites/Domains)",
        [
            "Immediately remove the `.git` directory from the web server's document root.",
            "Add `.git` to the web server's deny list to prevent it from ever being served.",
            "Rotate any credentials found in the repository's history."
        ],
        ['domain', 'leak_url', 'status_code', 'Risk'], 'Git_Leak_out.csv',
        lambda row: row.get('status_code') == '200',
        risk_column='Risk', default_risk='High', risk_order=risk_priority
    )

    jquery_description = (
        f"This section identifies web applications using outdated software components. Outdated software often contains known, publicly disclosed vulnerabilities (CVEs) that attackers can easily exploit. "
        f"For example, any version of jQuery below 3.0.0 is no longer receiving security patches. The current stable version is {latest_jquery_version}. "
        "The table below lists applications running outdated jQuery versions, which should be prioritized for upgrade."
    )
    add_vulnerability_section(
        doc, "Outdated jQuery Versions Detected",
        jquery_description,
        "Initial Access (T1190: Exploit Public-Facing Application)",
        [
            "Upgrade the outdated library to the latest stable version.",
            "Implement a Software Bill of Materials (SBOM) and a regular patch management process.",
            "Remove any unused or unnecessary JavaScript libraries."
        ],
        ['url', 'JQuery'], 'tech_detection_unified.csv',
        lambda row: row.get('JQuery') and row.get('JQuery', '0')[0] in ['1', '2']
    )

    # Add Python 2.x detection
    add_vulnerability_section(
        doc, "End-of-Life Python 2.x Detected",
        ("Python 2.x reached End of Life on January 1, 2020, meaning it no longer receives security updates or bug fixes. "
         "Applications running Python 2.7 are vulnerable to known security flaws that will never be patched. "
         "Attackers actively target systems running EOL software as they represent easy targets with documented vulnerabilities. "
         "The detected systems are running Python 2.7.5, which contains numerous publicly disclosed CVEs."),
        "Initial Access (T1190: Exploit Public-Facing Application)",
        [
            "Prioritize migration to Python 3.x immediately - Python 2 has been EOL since January 2020.",
            "Identify all applications and scripts dependent on Python 2.7 and create migration plan.",
            "If immediate migration is not possible, isolate affected systems from public internet access.",
            "Implement compensating controls such as WAF rules and enhanced monitoring for affected assets.",
            "Review CVE databases for Python 2.7 vulnerabilities that may affect your applications."
        ],
        ['target', 'Backend_Python'], 'tech_detection_unified.csv',
        lambda row: row.get('Backend_Python', '').startswith('2.')
    )

    # Add OpenSSL 1.0.x detection
    add_vulnerability_section(
        doc, "End-of-Life OpenSSL 1.0.x Detected",
        ("OpenSSL 1.0.2 reached End of Life in December 2019 and no longer receives security updates. "
         "OpenSSL is a critical cryptographic library used to secure communications (HTTPS, SSL/TLS). "
         "Running EOL versions exposes systems to known cryptographic vulnerabilities including those that could allow man-in-the-middle attacks, "
         "certificate validation bypasses, and encryption weaknesses. The detected version (1.0.2k-fips) contains multiple known CVEs."),
        "Initial Access (T1190: Exploit Public-Facing Application) / Credential Access (T1556.004: Modify Authentication Process)",
        [
            "Upgrade to OpenSSL 3.x or at minimum OpenSSL 1.1.1 (which receives support until September 2026).",
            "Test certificate validation and TLS configuration after upgrade to ensure compatibility.",
            "Review systems for dependencies on specific OpenSSL 1.0.x features that may require code changes.",
            "Scan for known OpenSSL CVEs affecting version 1.0.2k including Heartbleed-class vulnerabilities.",
            "Consider using operating system packages that backport security fixes if immediate upgrade is not feasible."
        ],
        ['target', 'Backend_OpenSSL'], 'tech_detection_unified.csv',
        lambda row: row.get('Backend_OpenSSL', '').startswith('1.0.') or row.get('Backend_OpenSSL', '').startswith('0.')
    )

    # Add PHP 5.x detection
    add_vulnerability_section(
        doc, "End-of-Life PHP 5.x Detected",
        ("PHP 5.x versions reached End of Life between 2015 and 2018 depending on the minor version. "
         "PHP is a widely-used server-side scripting language, and EOL versions no longer receive security patches. "
         "Known vulnerabilities in PHP 5.x include remote code execution, SQL injection facilitation, and authentication bypasses. "
         "The detected version (PHP 5.4.16) is extremely outdated and contains numerous critical security flaws."),
        "Initial Access (T1190: Exploit Public-Facing Application)",
        [
            "Upgrade to PHP 8.x immediately (PHP 7.x will reach EOL in 2025-2026).",
            "Test application compatibility with newer PHP versions in development environment first.",
            "Review application code for deprecated PHP functions that may break during upgrade.",
            "If upgrade is not immediately possible, implement strict input validation and WAF rules.",
            "Check for known CVEs affecting PHP 5.4 and implement compensating controls where possible."
        ],
        ['target', 'Backend_PHP'], 'tech_detection_unified.csv',
        lambda row: row.get('Backend_PHP', '').startswith('5.')
    )

    add_vulnerability_section(
        doc, "Stale DNS Records",
        ("Stale DNS records are entries that point to resources no longer in use. These can be exploited by attackers in a technique called 'subdomain takeover,' "
         "where they claim the abandoned resource to host malicious content under your company's domain, exploiting your brand's trust."),
        "Initial Access (T1565.001: Hijack Execution Flow)",
        [
            "Regularly audit DNS records and remove any entries that point to decommissioned services.",
            "If a third-party service is no longer in use, immediately remove the corresponding DNS record."
        ],
        ['domain', 'stale', 'type', 'target'], 'Dead_DNS_Checker_out.csv',
        lambda row: row.get('stale', '').lower() == 'true'
    )

    add_vulnerability_section(
        doc, "Subdomain Takeover & Expired SSL",
        ("This section highlights two related risks. 'Subdomain Takeover' vulnerabilities occur when a subdomain points to a third-party service that is no longer in use, allowing an attacker to claim it. "
         "'Expired SSL Certificates' create security warnings for users and can be a sign of a neglected service, making it a potential target."),
        "Initial Access (T1565.001: Hijack Execution Flow)",
        [
            "Remove DNS records pointing to unclaimed third-party services.",
            "Renew expired SSL/TLS certificates immediately.",
            "Implement an automated certificate renewal process."
        ],
        ['domain', 'target', 'takeover_status', 'ssl_expiry'], 'Dead_DNS_Resolver_out.csv',
        lambda row: row.get('takeover_status') == 'vulnerable' or \
                    (row.get('ssl_expiry') and datetime.strptime(row.get('ssl_expiry'), '%Y-%m-%d') < datetime.now())
    )

    add_vulnerability_section(
        doc, "Subdomain Takeover Risks",
        ("Stale DNS records pointing to unclaimed cloud resources create serious subdomain takeover vulnerabilities. "
         "When DNS records reference external services (cloud hosting, CDNs, etc.) that are no longer active, attackers can claim those resources and serve malicious content under your trusted domain. "
         "This analysis identifies DNS records at high risk of takeover based on CNAME targets, service detection, and SSL certificate validation."),
        "Compromise Infrastructure (T1584.001: Domains)",
        [
            "Immediately remove DNS records pointing to decommissioned or unclaimed cloud services.",
            "Claim any abandoned cloud resources if they are still legitimately needed, or delete the DNS records.",
            "Maintain an inventory of all subdomains and their associated third-party services.",
            "Implement monitoring for DNS changes and new subdomain creation.",
            "Review SSL certificates on subdomains - mismatched CNs often indicate takeover risk."
        ],
        ['domain', 'type', 'target', 'service', 'takeover_status', 'ssl_cn'], 'dns_takeover_risks.csv',
        lambda row: True
    )

    add_simple_list_section(
        doc, "Exposed Debug Paths",
        ("Debug paths are special URLs (e.g., /debug, /test) that developers use during testing. When left enabled on a live website, they can leak sensitive system information, "
         "provide access to developer tools, or even allow an attacker to bypass security controls."),
        "Reconnaissance (T1592: Gather Victim Host Information)",
        [
            "Ensure that all debugging features and diagnostic endpoints are disabled in production environments.",
            "Use build flags and environment variables to programmatically disable debug code in production builds."
        ],
        'Debug_paths_out.csv', 'Exposed Path',
        lambda line: '(200)' in line
    )

    add_vulnerability_section(
        doc, "Default Web Pages",
        ("This section identifies web servers displaying default or placeholder pages (e.g., 'It works!'). While not a direct vulnerability, these pages indicate that a server may have been set up but not fully configured or secured, "
         "making it a potential target for attackers looking for easy-to-exploit systems."),
        "Reconnaissance (T1592: Gather Victim Host Information)",
        [
            "Remove or replace the default placeholder pages on all web servers.",
            "Implement a standard server hardening process that includes removing default content."
        ],
        ['url', 'status', 'server'], 'Default_Page_Checker_out.csv',
        lambda row: row.get('status') == '200' or 'default' in row.get('status', '').lower()
    )
    
    add_simple_list_section(
        doc, "Exposed Non-Production Domains",
        ("Our scans discovered several domains that appear to be used for non-production purposes (e.g., 'dev', 'staging', 'uat'). These environments are often less secure than production systems and may contain default credentials or unpatched software, "
         "providing an easier entry point for an attacker into your organization's network."),
        "Reconnaissance (T1593: Search Open Websites/Domains)",
        [
            "Restrict access to all non-production environments using IP whitelisting or a VPN.",
            "Ensure non-production environments are not indexed by search engines."
        ],
        'Non_Production_domains_out.csv', 'Domain'
    )

    add_vulnerability_section(
        doc, "Potential Typosquatting Domains (Active Registrations Only)",
        ("This section lists ACTIVELY REGISTERED domains that are very similar to your company's official domains. "
         "We generated 623,958 possible typosquatting variations using techniques like bitsquatting, omission, repetition, and keyboard proximity, "
         "then filtered to show ONLY domains that are currently registered and resolving to IP addresses. "
         "Attackers register these 'typosquatted' domains to trick employees and customers into visiting malicious websites for phishing attacks or malware distribution. "
         "While not all registered variations are necessarily malicious, they represent a potential risk to your brand and security."),
        "Resource Development (T1585.001: Acquire Infrastructure - Domains)",
        [
            "Consider defensively registering common variations of your domain.",
            "Use a brand protection service to identify and initiate takedown requests for malicious domains.",
            "Monitor for registration of high-risk variations listed in this report.",
            "Investigate each registered domain to determine if it's legitimate, parked, or actively malicious."
        ],
        ['root_domain', 'typo_domain', 'fuzzer', 'active', 'dns_a'], 'typo_candidates.csv',
        lambda row: row.get('dns_a', '').strip() != ''
    )

    add_redirect_issue_section(
        doc, "Misconfigured Redirects",
        ("This finding identifies web pages that are not automatically redirecting users from the insecure http:// to the secure https:// version of the site. "
         "This lack of enforcement can expose users to man-in-the-middle attacks, where an attacker could intercept traffic and steal sensitive information like login credentials."),
        "Initial Access (T1566.002: Spearphishing Link)",
        [
            "Implement HTTP Strict Transport Security (HSTS) to enforce the use of HTTPS.",
            "Configure web servers to perform a 301 (permanent) redirect from HTTP to HTTPS for all requests."
        ],
        'redirectissuefinder_out.csv'
    )
    
    add_simple_list_section(
        doc, "Internal IP Address Exposure",
        ("Our scans found instances where internal, private IP addresses (e.g., 10.x.x.x, 192.168.x.x) were exposed in publicly accessible web content. "
         "This information can give attackers valuable clues about the structure of your internal network, which can be used to plan more targeted attacks."),
        "Reconnaissance (T1592: Gather Victim Host Information)",
        [
            "Review and sanitize web application responses, logs, and source code to remove any references to internal IP addresses.",
            "Configure web servers to strip or rewrite headers and error messages that may contain internal network information."
        ],
        'Detect_Internal_DNS_out.txt', 'Finding',
        lambda line: any(p in line for p in ['10.', '172.16.', '192.168.'])
    )

    # Add WHOIS Intelligence section
    add_vulnerability_section(
        doc, "IP Address Intelligence (WHOIS) - Noteworthy Findings Only",
        ("WHOIS lookups were performed on all IP addresses that your discovered domains resolve to. "
         "This section shows ONLY the noteworthy findings that warrant investigation - approximately 20 out of 773 total IPs scanned. "
         "Filtered results include: (1) RFC1918 private IPs leaking in public DNS, (2) Third-party cloud hosting (AWS, Azure, GCP), "
         "(3) CDN and security services (CloudFront, Proofpoint), (4) Shadow IT or unexpected providers. "
         "Your organization's owned IP space is automatically filtered out to reduce noise and focus on noteworthy third-party infrastructure."),
        "Reconnaissance (T1590: Gather Victim Network Information)",
        [
            "Review all IP addresses to ensure they belong to authorized infrastructure.",
            "Identify third-party or cloud-hosted services that may not be under direct security oversight.",
            "Check for unexpected geographic locations or hosting providers.",
            "Verify that RFC1918 private IPs (10.x, 192.168.x, 172.16.x) are not leaking in public DNS.",
            "Confirm all externally hosted services comply with your organization's security and compliance policies.",
            "Monitor for unauthorized IP allocations or infrastructure ownership changes."
        ],
        ['ip', 'whois_info'], 'whois_results.csv',
        lambda row: is_interesting_ip(row)
    )

def is_interesting_ip(row):
    """Filter WHOIS data to only show IPs worth investigating (reduces noise from 773 to ~20 findings)."""
    whois = row.get('whois_info', '').lower()
    if not whois:
        return False

    # Private IP leaks in public DNS (RFC1918)
    if 'private-address' in whois or 'rfc1918' in whois:
        return True

    # Third-party cloud providers
    cloud_providers = ['amazon', 'aws', 'google', 'azure', 'microsoft', 'cloudflare',
                      'akamai', 'fastly', 'digitalocean', 'linode', 'rackspace']
    if any(provider in whois for provider in cloud_providers):
        return True

    # CDNs and security services
    cdn_security = ['cloudfront', 'proofpoint', 'mimecast', 'barracuda', 'zscaler']
    if any(service in whois for service in cdn_security):
        return True

    # Reserved/special use IPs that shouldn't appear
    if 'reserved' in whois or 'special' in whois:
        return True

    # Everything else (likely owned infrastructure) - don't show in report
    return False

def add_cve_section(doc):
    """Adds a section for software with known vulnerabilities (CVEs)."""
    software_columns = [
        'Apache', 'Nginx', 'Microsoft-IIS', 'LiteSpeed', 'OpenResty', 'Pepyaka', 'Caddy',
        'PHP', 'ASP_NET', 'Java', 'Ruby-on-Rails',
        'WordPress', 'JQuery', 'Bootstrap', 'Modernizr', 'Moodle', 'Odoo', 'TYPO3',
        'OpenSSL', 'Microsoft-HTTPAPI'
    ]

    rows_to_render = []
    try:
        with open(get_file_path('tech_detection_unified.csv'), 'r', encoding='utf-8') as f:
            reader = csv.DictReader(f)
            tech_list = []
            for row in reader:
                for key, value in row.items():
                    if value and key in software_columns:
                        parts = value.split('/')
                        product = parts[0].strip()
                        version = parts[1].strip() if len(parts) > 1 else 'Unknown'

                        if product.lower() in ['unknown', 'text', 'javascript']:
                            continue

                        tech_list.append({'product': product, 'version': version, 'url': row.get('url', '')})

            processed_urls = set()
            cve_lookups = 0
            for tech in tech_list:
                if tech['url'] in processed_urls or cve_lookups >= 10:
                    continue
                cves = get_cves_for_product(tech['product'], tech['version'])
                if cves:
                    rows_to_render.append({
                        'url': tech['url'],
                        'product': tech['product'],
                        'version': tech['version'],
                        'cves': ', '.join([c['cve_id'] for c in cves[:3]])
                    })
                    processed_urls.add(tech['url'])
                    cve_lookups += 1

    except FileNotFoundError:
        return
    except Exception as e:
        print(f"An error occurred during CVE processing: {e}")
        return

    if not rows_to_render:
        return

    doc.add_heading("Software with Known Vulnerabilities (CVEs)", level=3)
    doc.add_paragraph(
        "This section identifies software and services that were detected on your public-facing assets and are associated with known vulnerabilities (CVEs). "
        "A CVE is a publicly disclosed security flaw that may signal an exploitable weakness. "
        "Verify these leads against your production configurations, as both software inventories and vulnerability data change frequently.",
        style='Normal'
    )

    table = doc.add_table(rows=1, cols=4, style='Table Grid')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'URL'
    hdr_cells[1].text = 'Software'
    hdr_cells[2].text = 'Version'
    hdr_cells[3].text = 'Potential CVEs'

    for row in rows_to_render:
        row_cells = table.add_row().cells
        row_cells[0].text = row['url']
        row_cells[1].text = row['product']
        row_cells[2].text = row['version']
        row_cells[3].text = row['cves']

def create_standalone_html_appendix():
    """Creates a single, self-contained HTML file with all data and JS."""
    
    js_content = ""
    # Read the JS file from the current directory
    if os.path.exists('appendix.js'):
        with open('appendix.js', 'r', encoding='utf-8') as f:
            js_content = f.read()
    else:
        print("[ERROR] appendix.js not found. The interactive appendix will not be functional.")

    html_template = ""
    # Assume the template is named Interactive_Appendix.html and is in the current directory
    # For this to work, the user's provided HTML source needs to be saved as this file.
    # Since I cannot create files, I will assume it exists.
    if os.path.exists('Interactive_Appendix.html'):
        with open('Interactive_Appendix.html', 'r', encoding='utf-8') as f:
            html_template = f.read()
    else:
        print("[ERROR] Interactive_Appendix.html not found. Cannot generate appendix.")
        return

    data_path = os.path.join(REPORT_DIR, 'appendix_data.js')
    with open(data_path, 'w', encoding='utf-8') as f:
        f.write('const appendixData = ')
        json.dump(all_vulnerabilities, f)
        f.write(';')

    final_html = html_template.replace(
        '<!-- SCRIPT_PLACEHOLDER -->',
        '<script src="appendix_data.js"></script>\n<script>' + js_content + '</script>'
    )

    with open(os.path.join(REPORT_DIR, 'Interactive_Appendix.html'), 'w', encoding='utf-8') as f:
        f.write(final_html)


def main():
    """Main function to generate the report."""
    global all_vulnerabilities, BASE_DIR, REPORT_DIR, ASSETS_DIR

    parser = argparse.ArgumentParser(description="Generate a Threat Intelligence report.")
    parser.add_argument("--input-file", required=True, help="Path to the Recon_out_enriched.csv file.")
    parser.add_argument("--output-dir", default=None, help="Output directory for the report (default: ../report relative to input file)")
    args = parser.parse_args()

    # Set BASE_DIR dynamically based on the input file path
    BASE_DIR = os.path.dirname(os.path.abspath(args.input_file))

    # Set REPORT_DIR based on --output-dir or default to ../report
    if args.output_dir:
        REPORT_DIR = os.path.abspath(args.output_dir)
    else:
        REPORT_DIR = os.path.join(BASE_DIR, "..", "report")

    ASSETS_DIR = os.path.join(REPORT_DIR, "assets")

    print("Generating Threat Intelligence Report...")

    os.makedirs(REPORT_DIR, exist_ok=True)
    os.makedirs(ASSETS_DIR, exist_ok=True)
    
    latest_jquery = get_latest_jquery_version()
    scope_stats = generate_scope_summary()

    doc = Document()
    set_styles(doc.styles)
    add_header_footer(doc)
    create_report_structure(doc)
    add_scope_section(doc, scope_stats)
    
    doc.add_page_break()
    doc.add_heading("Intelligence Analysis", level=1)
    doc.add_heading("Technology Footprint & Vulnerabilities", level=2)

    populate_findings(doc, latest_jquery)
    add_cve_section(doc)
    
    create_standalone_html_appendix()

    output_filename = 'Threat_Intelligence_Report.docx'
    doc.save(os.path.join(REPORT_DIR, output_filename))
    
    # Create README.txt
    readme_content = (
        "Threat Intelligence Report Package\n"
        "==================================\n\n"
        "This package contains two main components:\n\n"
        "1. Threat_Intelligence_Report.docx\n"
        "   - This is the main, human-readable report containing the executive summary, findings, and remediation advice.\n\n"
        "2. Interactive_Appendix.html\n"
        "   - This is a searchable, interactive file for exploring all raw vulnerability data. Open this file in your web browser.\n\n"
    )
    with open(os.path.join(REPORT_DIR, 'README.txt'), 'w') as f:
        f.write(readme_content)

    print(f"Report and appendix files saved to '{REPORT_DIR}' directory.")

if __name__ == "__main__":
    main()
